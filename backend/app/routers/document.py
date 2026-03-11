from fastapi import APIRouter, HTTPException, status, Header, UploadFile, File
from fastapi.responses import Response
from typing import Optional, List
from pydantic import BaseModel
import re
import io
import os
import tempfile
import subprocess
import pdfplumber
from docx import Document
from ..routers.profile import get_current_user
from ..services.ai_service import ai_service
from ..database import get_supabase_admin
import logging
import json

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/document", tags=["Document Parser"])

# Common patterns for CV parsing
EMAIL_PATTERN = r'[\w\.-]+@[\w\.-]+\.\w+'
PHONE_PATTERN = r'(?:\+?\d{1,4}[\s\-\.]?)?\(?\d{1,4}\)?[\s\-\.]?\d{1,4}[\s\-\.]?\d{1,9}'
LINKEDIN_PATTERN = r'(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9\-_]+)|linkedin[:\s]+([a-zA-Z0-9\-_]+)'
GITHUB_PATTERN = r'(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9\-_]+)|github[:\s]+([a-zA-Z0-9\-_]+)'
WEBSITE_PATTERN = r'(?:portfolio|website|web)[:\s]*\s*(https?://[^\s]+)|(?<![@\w])(https?://(?!linkedin|github)[^\s]+)'

# Section headers to identify different parts of CV
SECTION_HEADERS = {
    'experience': ['experience', 'work experience', 'employment', 'employment history', 'work history', 'professional experience', 'career history', 'professional background', 'positions held'],
    'education': ['education', 'academic background', 'academic', 'qualification', 'qualifications', 'degree', 'degrees', 'university', 'college', 'schooling', 'educational background'],
    'skills': ['skills', 'technical skills', 'technologies', 'competencies', 'expertise', 'proficiencies', 'core competencies', 'technical expertise', 'areas of expertise', 'key skills', 'professional skills'],
    'projects': ['projects', 'personal projects', 'portfolio', 'side projects', 'key projects', 'notable projects', 'selected projects'],
    'certifications': ['certifications', 'certificates', 'licenses', 'credentials', 'professional certifications', 'training', 'courses'],
    'languages': ['languages', 'language proficiency', 'language skills', 'spoken languages'],
    'summary': ['summary', 'objective', 'profile', 'about me', 'about', 'professional summary', 'career objective', 'executive summary', 'professional profile', 'career summary', 'overview'],
}

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file with improved accuracy"""
    text = ""
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            # Try to extract text with better settings for resume formatting
            page_text = page.extract_text(
                x_tolerance=3,
                y_tolerance=3,
                layout=True,
                x_density=7.25,
                y_density=13
            )
            if page_text:
                text += page_text + "\n\n"
            
            # Also try to extract tables for structured data
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        row_text = ' | '.join([cell.strip() if cell else '' for cell in row if cell])
                        if row_text.strip():
                            text += row_text + "\n"
            
            # Extract words with their positions for better layout understanding
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False
            )
            
            # Group words by their vertical position (lines)
            if words:
                lines_by_y = {}
                for word in words:
                    y_key = round(word['top'] / 10) * 10  # Group by ~10 pixel rows
                    if y_key not in lines_by_y:
                        lines_by_y[y_key] = []
                    lines_by_y[y_key].append(word)
                
                # Sort lines and extract text
                for y_key in sorted(lines_by_y.keys()):
                    line_words = sorted(lines_by_y[y_key], key=lambda w: w['x0'])
                    line_text = ' '.join([w['text'] for w in line_words])
                    # Only add if significantly different from layout extraction
                    if line_text.strip() and line_text not in text:
                        text += line_text + "\n"
    
    # Clean up the text
    text = re.sub(r'\n{4,}', '\n\n\n', text)  # Remove excessive newlines
    text = re.sub(r' {3,}', '  ', text)  # Reduce excessive spaces
    text = re.sub(r'(\S)\s*\n\s*(\S)', r'\1 \2', text)  # Join broken lines
    
    return text.strip()

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    doc = Document(io.BytesIO(file_content))
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    return text

def find_section(text: str, section_keywords: list) -> tuple[int, int]:
    """Find the start position of a section"""
    text_lower = text.lower()
    for keyword in section_keywords:
        # Look for section header patterns
        patterns = [
            rf'\n\s*{keyword}\s*\n',
            rf'\n\s*{keyword}\s*:',
            rf'\n\s*{keyword.upper()}\s*\n',
            rf'^{keyword}\s*:',
        ]
        for pattern in patterns:
            match = re.search(pattern, text_lower, re.MULTILINE | re.IGNORECASE)
            if match:
                return match.start(), match.end()
    return -1, -1

def extract_section_content(text: str, section_name: str, all_sections: dict) -> str:
    """Extract content for a specific section"""
    section_keywords = SECTION_HEADERS.get(section_name, [])
    start, header_end = find_section(text, section_keywords)
    
    if start == -1:
        return ""
    
    # Find where this section ends (when another section begins)
    next_section_start = len(text)
    for other_section, other_keywords in SECTION_HEADERS.items():
        if other_section == section_name:
            continue
        other_start, _ = find_section(text[header_end:], other_keywords)
        if other_start != -1:
            actual_start = header_end + other_start
            if actual_start < next_section_start:
                next_section_start = actual_start
    
    return text[header_end:next_section_start].strip()

def parse_experience(content: str) -> list:
    """Parse experience section into structured data with improved accuracy"""
    experiences = []
    if not content:
        return experiences
    
    lines = content.split('\n')
    current_exp = None
    
    # Enhanced date patterns
    date_patterns = [
        # Full month names: January 2020 - Present
        r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s*\d{4})\s*[-–—to]+\s*((?:January|February|March|April|May|June|July|August|September|October|November|December)\s*\d{4}|[Pp]resent|[Cc]urrent|[Oo]ngoing|[Nn]ow)',
        # Abbreviated months: Jan 2020 - Present
        r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–—to]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|[Pp]resent|[Cc]urrent|[Oo]ngoing|[Nn]ow)',
        # Numeric: 01/2020 - 12/2023 or 2020-2023
        r'(\d{1,2}[/\-\.]\d{4}|\d{4})\s*[-–—to]+\s*(\d{1,2}[/\-\.]\d{4}|\d{4}|[Pp]resent|[Cc]urrent|[Oo]ngoing|[Nn]ow)',
        # Year only in parentheses: (2020 - 2023)
        r'\((\d{4})\s*[-–—to]+\s*(\d{4}|[Pp]resent|[Cc]urrent)\)',
    ]
    
    def find_date_match(line):
        for pattern in date_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                return match
        return None
    
    def normalize_date(date_str):
        """Convert date string to YYYY-MM format"""
        if not date_str:
            return ''
        date_str = date_str.strip()
        
        if date_str.lower() in ['present', 'current', 'ongoing', 'now']:
            return ''
        
        # Try to extract year and month
        month_map = {
            'january': '01', 'jan': '01', 'february': '02', 'feb': '02',
            'march': '03', 'mar': '03', 'april': '04', 'apr': '04',
            'may': '05', 'june': '06', 'jun': '06', 'july': '07', 'jul': '07',
            'august': '08', 'aug': '08', 'september': '09', 'sep': '09', 'sept': '09',
            'october': '10', 'oct': '10', 'november': '11', 'nov': '11',
            'december': '12', 'dec': '12'
        }
        
        # Match month name + year
        month_year = re.match(r'([A-Za-z]+)\s*(\d{4})', date_str)
        if month_year:
            month = month_map.get(month_year.group(1).lower(), '01')
            return f"{month_year.group(2)}-{month}"
        
        # Match MM/YYYY or MM-YYYY
        num_date = re.match(r'(\d{1,2})[/\-\.](\d{4})', date_str)
        if num_date:
            return f"{num_date.group(2)}-{num_date.group(1).zfill(2)}"
        
        # Year only
        year_only = re.match(r'(\d{4})', date_str)
        if year_only:
            return f"{year_only.group(1)}-01"
        
        return date_str
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        
        # Check if line contains a date range
        date_match = find_date_match(line)
        
        # Detect new experience entry
        is_new_entry = False
        if date_match:
            is_new_entry = True
        elif len(line) < 120 and not line.startswith(('•', '-', '*', '●', '○', '►', '▪')):
            # Could be title/company without date on same line
            # Check if next lines have dates or bullets
            has_context = any(find_date_match(lines[j]) or lines[j].strip().startswith(('•', '-', '*')) 
                            for j in range(i+1, min(i+4, len(lines))))
            if has_context:
                is_new_entry = True
        
        if is_new_entry:
            # Save previous experience
            if current_exp and (current_exp.get('title') or current_exp.get('company')):
                experiences.append(current_exp)
            
            current_exp = {
                'title': '',
                'company': '',
                'location': '',
                'start_date': '',
                'end_date': '',
                'current': False,
                'description': '',
                'achievements': []
            }
            
            if date_match:
                current_exp['start_date'] = normalize_date(date_match.group(1))
                end_date = date_match.group(2)
                if end_date.lower() in ['present', 'current', 'ongoing', 'now']:
                    current_exp['current'] = True
                    current_exp['end_date'] = ''
                else:
                    current_exp['end_date'] = normalize_date(end_date)
                
                # Remove date from line
                line_without_date = line[:date_match.start()] + line[date_match.end():]
            else:
                line_without_date = line
            
            line_without_date = line_without_date.strip(' -–—|,')
            
            # Parse title, company, location from the remaining text
            # Common patterns: "Title at Company" or "Title | Company | Location" or "Title - Company"
            separators = [
                (r'\s+at\s+', ['title', 'company']),
                (r'\s*\|\s*', ['title', 'company', 'location']),
                (r'\s*[-–—]\s*(?![Pp]resent)', ['title', 'company']),
                (r'\s*,\s*', ['title', 'company', 'location']),
            ]
            
            parsed = False
            for sep_pattern, fields in separators:
                parts = re.split(sep_pattern, line_without_date, maxsplit=len(fields)-1)
                if len(parts) >= 2:
                    for idx, field in enumerate(fields):
                        if idx < len(parts) and parts[idx].strip():
                            current_exp[field] = parts[idx].strip()
                    parsed = True
                    break
            
            if not parsed and line_without_date:
                current_exp['title'] = line_without_date
        
        elif current_exp and line.startswith(('•', '-', '*', '●', '○', '►', '▪', '→')):
            # This is a bullet point (achievement)
            achievement = re.sub(r'^[•\-*●○►▪→]\s*', '', line).strip()
            if achievement and len(achievement) > 5:
                current_exp['achievements'].append(achievement)
        
        elif current_exp and not current_exp.get('company') and len(line) < 80:
            # This might be the company name on a separate line
            if not line.startswith(('•', '-', '*')) and not find_date_match(line):
                if not current_exp['company']:
                    current_exp['company'] = line
                elif not current_exp['location'] and len(line) < 50:
                    current_exp['location'] = line
    
    # Don't forget the last experience
    if current_exp and (current_exp.get('title') or current_exp.get('company')):
        experiences.append(current_exp)
    
    # Generate descriptions from achievements
    for exp in experiences:
        if exp['achievements'] and not exp['description']:
            exp['description'] = '\n'.join(['• ' + a for a in exp['achievements'][:5]])
    
    return experiences[:15]  # Limit to 15 experiences

def parse_education(content: str) -> list:
    """Parse education section into structured data"""
    education = []
    if not content:
        return education
    
    lines = content.split('\n')
    current_edu = None
    
    # Common degree patterns
    degree_patterns = [
        r'(Bachelor|Master|PhD|Ph\.D|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|B\.?E\.?|M\.?E\.?|B\.?Tech|M\.?Tech|MBA|BBA)',
        r'(Diploma|Certificate|Associate)',
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if this line contains a degree
        is_degree_line = any(re.search(p, line, re.IGNORECASE) for p in degree_patterns)
        
        if is_degree_line or (len(line) < 150 and not line.startswith(('•', '-', '*'))):
            if current_edu and (current_edu.get('degree') or current_edu.get('institution')):
                education.append(current_edu)
            
            current_edu = {
                'degree': '',
                'field_of_study': '',
                'institution': '',
                'location': '',
                'start_date': '',
                'end_date': '',
                'gpa': '',
                'description': '',
                'achievements': []
            }
            
            # Extract dates
            date_pattern = r'(\d{4})\s*[-–—to]+\s*(\d{4}|[Pp]resent)'
            date_match = re.search(date_pattern, line)
            if date_match:
                current_edu['start_date'] = date_match.group(1)
                current_edu['end_date'] = date_match.group(2)
            
            # Extract GPA
            gpa_match = re.search(r'GPA[:\s]*([0-9.]+)', line, re.IGNORECASE)
            if gpa_match:
                current_edu['gpa'] = gpa_match.group(1)
            
            # Try to identify degree and institution
            for pattern in degree_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    current_edu['degree'] = line
                    break
    
    if current_edu and (current_edu.get('degree') or current_edu.get('institution')):
        education.append(current_edu)
    
    return education[:5]  # Limit to 5 education entries

def parse_skills(content: str) -> list:
    """Parse skills section into grouped skill data.
    
    Recognises two common formats:
    1. Category-based (tabular or colon-separated):
       Languages: Python, JavaScript, TypeScript
       AI / ML: LangChain, Transformers, ...
    2. Flat list (comma/bullet separated) → placed under 'Technical'.
    
    Returns list of {category: str, items: [str]}
    """
    if not content:
        return []

    lines = content.strip().split('\n')
    groups: list[dict] = []
    current_category = None
    current_items: list[str] = []

    # Pattern: "Category:" or "Category  item1, item2" (tabular with 2+ spaces or | separator)
    category_pattern = re.compile(
        r'^([A-Za-z][A-Za-z &/\-]+?)\s*[:|\t]\s*(.+)$'
    )
    # Also match tabular format: "Category    item1, item2" (2+ spaces between)
    tabular_pattern = re.compile(
        r'^([A-Za-z][A-Za-z &/\-]+?)\s{2,}(.+)$'
    )

    skip_words = {'and', 'or', 'the', 'a', 'an', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'etc', 'including'}

    def split_items(text: str) -> list[str]:
        """Split a comma/semicolon/bullet separated skill string into clean items."""
        parts = re.split(r'[,;•●○\|]+', text)
        items = []
        for p in parts:
            p = p.strip().strip('-').strip('*').strip()
            if p and len(p) > 1 and len(p) < 80 and p.lower() not in skip_words:
                items.append(p)
        return items

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Try category: items pattern
        match = category_pattern.match(line) or tabular_pattern.match(line)
        if match:
            # Save previous category if exists
            if current_category and current_items:
                groups.append({'category': current_category, 'items': current_items})
            
            current_category = match.group(1).strip()
            current_items = split_items(match.group(2))
        elif current_category:
            # Continuation line for current category
            new_items = split_items(line)
            if new_items:
                current_items.extend(new_items)
        else:
            # No category detected yet — flat items
            new_items = split_items(line)
            if new_items:
                current_items.extend(new_items)

    # Save last group
    if current_category and current_items:
        groups.append({'category': current_category, 'items': current_items})
    elif current_items:
        # All flat items, no categories detected
        groups.append({'category': 'Technical', 'items': current_items})

    # If nothing was parsed as grouped, fall back to flat
    if not groups:
        all_items = split_items(content.replace('\n', ', '))
        if all_items:
            groups.append({'category': 'Technical', 'items': all_items[:40]})

    return groups

def parse_projects(content: str) -> list:
    """Parse projects section into structured data"""
    projects = []
    if not content:
        return projects
    
    lines = content.split('\n')
    current_project = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Project names are usually short and not bullet points
        if len(line) < 100 and not line.startswith(('•', '-', '*', '●')):
            if current_project and current_project.get('name'):
                projects.append(current_project)
            
            current_project = {
                'name': line,
                'description': '',
                'technologies': [],
                'url': '',
                'github_url': '',
                'highlights': []
            }
            
            # Extract URLs
            url_match = re.search(r'https?://[^\s]+', line)
            if url_match:
                url = url_match.group()
                if 'github' in url.lower():
                    current_project['github_url'] = url
                else:
                    current_project['url'] = url
        
        elif current_project and line.startswith(('•', '-', '*', '●')):
            desc = line.lstrip('•-*●○ ').strip()
            if desc:
                current_project['highlights'].append(desc)
                if current_project['description']:
                    current_project['description'] += ' ' + desc
                else:
                    current_project['description'] = desc
    
    if current_project and current_project.get('name'):
        projects.append(current_project)
    
    return projects[:10]

def extract_personal_info(text: str) -> dict:
    """Extract personal information from CV text"""
    personal_info = {
        'full_name': '',
        'email': '',
        'phone': '',
        'address': '',
        'city': '',
        'linkedin_url': '',
        'github_url': '',
        'portfolio_url': '',
        'summary': ''
    }
    
    # Extract email
    email_match = re.search(EMAIL_PATTERN, text)
    if email_match:
        personal_info['email'] = email_match.group()
    
    # Extract phone - improved pattern
    phone_matches = re.findall(PHONE_PATTERN, text)
    if phone_matches:
        # Get the longest valid phone number
        valid_phones = [p.strip() for p in phone_matches if len(re.sub(r'\D', '', p)) >= 10]
        if valid_phones:
            personal_info['phone'] = valid_phones[0]
    
    # Extract LinkedIn - improved pattern
    linkedin_match = re.search(LINKEDIN_PATTERN, text, re.IGNORECASE)
    if linkedin_match:
        username = linkedin_match.group(1) or linkedin_match.group(2)
        if username:
            personal_info['linkedin_url'] = f"https://linkedin.com/in/{username}"
    
    # Try to find full LinkedIn URL
    full_linkedin = re.search(r'https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9\-_/]+', text, re.IGNORECASE)
    if full_linkedin:
        personal_info['linkedin_url'] = full_linkedin.group().rstrip('/')
    
    # Extract GitHub - improved pattern
    github_match = re.search(GITHUB_PATTERN, text, re.IGNORECASE)
    if github_match:
        username = github_match.group(1) or github_match.group(2)
        if username:
            personal_info['github_url'] = f"https://github.com/{username}"
    
    # Try to find full GitHub URL
    full_github = re.search(r'https?://(?:www\.)?github\.com/[a-zA-Z0-9\-_]+', text, re.IGNORECASE)
    if full_github:
        personal_info['github_url'] = full_github.group().rstrip('/')
    
    # Extract portfolio/website URL
    website_match = re.search(WEBSITE_PATTERN, text, re.IGNORECASE)
    if website_match:
        url = website_match.group(1) or website_match.group(2)
        if url and 'linkedin' not in url.lower() and 'github' not in url.lower():
            personal_info['portfolio_url'] = url.rstrip('/')
    
    # Extract name - improved algorithm
    lines = text.split('\n')
    first_lines = [l.strip() for l in lines[:15] if l.strip()]  # Check first 15 non-empty lines
    
    for line in first_lines:
        # Skip lines with email, phone, URLs, or section headers
        if re.search(r'[@|•\-\d{4}]|http|www\.|\.com|experience|education|skills|summary|objective', line, re.IGNORECASE):
            continue
        
        # Name is usually 2-4 words, all capitalized or title case
        if len(line) < 60 and len(line) > 2:
            words = line.split()
            if 2 <= len(words) <= 5:
                # Check if it looks like a name (mostly letters, possibly with . or ,)
                cleaned = re.sub(r'[,.\s]', '', line)
                if cleaned.isalpha() or re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$', line):
                    personal_info['full_name'] = line.strip()
                    break
    
    # If no name found, try first line that looks like a name
    if not personal_info['full_name']:
        for line in first_lines:
            if re.match(r'^[A-Z][A-Za-z\s\.]+$', line) and len(line) > 3 and len(line) < 50:
                if not any(header in line.lower() for headers in SECTION_HEADERS.values() for header in headers):
                    personal_info['full_name'] = line.strip()
                    break
    
    # Extract location/address - look for city, state patterns
    location_patterns = [
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),?\s*([A-Z]{2})\s*(\d{5})?',  # City, ST ZIP
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),?\s*(India|USA|UK|Canada|Australia|Germany)',  # City, Country
        r'📍\s*([^\n]+)',  # Location with emoji
        r'(?:location|address|city)[:\s]+([^\n]+)',  # Labeled location
    ]
    
    for pattern in location_patterns:
        match = re.search(pattern, text[:1000], re.IGNORECASE)
        if match:
            personal_info['address'] = match.group(0).strip()
            if match.lastindex and match.group(1):
                personal_info['city'] = match.group(1).strip()
            break
    
    # Extract summary
    summary_content = extract_section_content(text, 'summary', SECTION_HEADERS)
    if summary_content:
        # Clean up summary
        summary = summary_content.strip()
        # Remove section headers if they got included
        for header in SECTION_HEADERS['summary']:
            summary = re.sub(rf'^{header}\s*[:.\-]?\s*', '', summary, flags=re.IGNORECASE)
        personal_info['summary'] = summary[:800]  # Limit summary length
    
    return personal_info

def parse_cv_document(text: str) -> dict:
    """Parse the full CV document and extract structured data"""
    result = {
        'personal_info': extract_personal_info(text),
        'experience': parse_experience(extract_section_content(text, 'experience', SECTION_HEADERS)),
        'education': parse_education(extract_section_content(text, 'education', SECTION_HEADERS)),
        'skills': parse_skills(extract_section_content(text, 'skills', SECTION_HEADERS)),
        'projects': parse_projects(extract_section_content(text, 'projects', SECTION_HEADERS)),
        'certifications': parse_certifications(extract_section_content(text, 'certifications', SECTION_HEADERS)),
        'languages': parse_languages(extract_section_content(text, 'languages', SECTION_HEADERS)),
        'raw_text': text[:10000]  # Include raw text for reference (limited)
    }
    
    return result

def parse_certifications(content: str) -> list:
    """Parse certifications section into structured data"""
    certifications = []
    if not content:
        return certifications
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 3:
            continue
        
        # Skip bullet point prefix
        if line.startswith(('•', '-', '*', '●', '○')):
            line = line.lstrip('•-*●○ ').strip()
        
        cert = {
            'name': '',
            'issuer': '',
            'date': '',
            'credential_id': '',
            'url': ''
        }
        
        # Try to extract date
        date_match = re.search(r'(\d{4}|\w+\s+\d{4})', line)
        if date_match:
            cert['date'] = date_match.group(1)
            line = line[:date_match.start()] + line[date_match.end():]
        
        # Try to split by common separators (-, by, from, |)
        parts = re.split(r'\s*[-–—|]\s*|\s+by\s+|\s+from\s+', line, maxsplit=1)
        cert['name'] = parts[0].strip()
        if len(parts) > 1:
            cert['issuer'] = parts[1].strip()
        
        # Extract credential ID if present
        cred_match = re.search(r'(?:ID|Credential)[:\s]*([A-Za-z0-9\-]+)', line, re.IGNORECASE)
        if cred_match:
            cert['credential_id'] = cred_match.group(1)
        
        if cert['name'] and len(cert['name']) > 2:
            certifications.append(cert)
    
    return certifications[:15]

def parse_languages(content: str) -> list:
    """Parse languages section into structured data"""
    languages = []
    if not content:
        return languages
    
    # Common proficiency levels
    proficiency_patterns = {
        'native': ['native', 'mother tongue', 'first language'],
        'fluent': ['fluent', 'bilingual', 'c2', 'c1'],
        'professional': ['professional', 'business', 'b2', 'advanced'],
        'intermediate': ['intermediate', 'b1', 'conversational'],
        'basic': ['basic', 'elementary', 'a1', 'a2', 'beginner']
    }
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 2:
            continue
        
        # Skip bullet points
        if line.startswith(('•', '-', '*', '●', '○')):
            line = line.lstrip('•-*●○ ').strip()
        
        lang = {
            'name': '',
            'proficiency': 'intermediate'  # Default
        }
        
        # Split by common separators
        parts = re.split(r'[:\-–—|()]', line)
        
        if parts:
            lang['name'] = parts[0].strip()
            
            # Look for proficiency in remaining text
            remaining = ' '.join(parts[1:]).lower() if len(parts) > 1 else ''
            
            for level, keywords in proficiency_patterns.items():
                if any(kw in remaining.lower() or kw in line.lower() for kw in keywords):
                    lang['proficiency'] = level
                    break
        
        if lang['name'] and len(lang['name']) > 1 and len(lang['name']) < 30:
            # Skip common non-language entries
            skip_words = ['skill', 'language', 'proficiency', 'level', 'spoken', 'written']
            if lang['name'].lower() not in skip_words:
                languages.append(lang)
    
    return languages[:10]

@router.post("/parse")
async def parse_document(
    file: UploadFile = File(...),
    authorization: Optional[str] = Header(None)
):
    """Parse a document (PDF or DOCX) and extract CV data"""
    # Auth is optional — allow unauthenticated access for ATS builder
    if authorization:
        try:
            user = get_current_user(authorization)
        except Exception:
            pass  # Continue without auth
    
    # Validate file type
    filename = file.filename.lower() if file.filename else ""
    if not filename.endswith(('.pdf', '.doc', '.docx')):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF, DOC, and DOCX files are supported"
        )
    
    try:
        # Read file content
        content = await file.read()
        
        if len(content) > 10 * 1024 * 1024:  # 10MB limit
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File size exceeds 10MB limit"
            )
        
        # Extract text based on file type
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(content)
        elif filename.endswith(('.doc', '.docx')):
            text = extract_text_from_docx(content)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unsupported file format"
            )
        
        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Could not extract text from the document. The file might be empty or contain only images."
            )
        
        logger.info(f"Extracted {len(text)} characters from document")
        
        # First, try rule-based parsing
        cv_data = parse_cv_document(text)
        
        # Try to enhance with AI if rule-based parsing seems incomplete
        try:
            if ai_service.client:
                ai_parsed = await parse_cv_with_ai(text)
                if ai_parsed:
                    # Merge AI results with rule-based results, preferring AI for complex fields
                    cv_data = merge_cv_data(cv_data, ai_parsed)
        except Exception as ai_error:
            logger.warning(f"AI parsing failed, using rule-based results: {str(ai_error)}")
        
        return {
            "success": True,
            "message": "Document parsed successfully",
            "data": cv_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error parsing document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to parse document: {str(e)}"
        )


async def parse_cv_with_ai(text: str) -> dict:
    """Use AI to parse CV text for better accuracy"""
    # Send up to 15000 chars to handle multi-page CVs
    cv_text = text[:15000]
    prompt = f"""Parse this CV/resume text and extract ALL the following information in JSON format.
Extract EVERY experience entry, project, certification, and skill category — do NOT skip any.

{{
    "personal_info": {{
        "full_name": "",
        "email": "",
        "phone": "",
        "address": "",
        "city": "",
        "linkedin_url": "",
        "github_url": "",
        "portfolio_url": "",
        "summary": "the professional summary or objective section text"
    }},
    "experience": [
        {{
            "title": "job title",
            "company": "company name",
            "location": "city or Remote",
            "start_date": "YYYY-MM format",
            "end_date": "YYYY-MM format or empty if current",
            "current": true/false,
            "description": "",
            "achievements": ["bullet point 1", "bullet point 2", "...extract ALL bullets"]
        }}
    ],
    "education": [
        {{
            "degree": "full degree name e.g. Bachelor of Technology (B.Tech) in Computer Science",
            "field_of_study": "specialization or major",
            "institution": "university/college name",
            "location": "city, state",
            "start_date": "YYYY-MM",
            "end_date": "YYYY-MM",
            "gpa": "e.g. 8.0 / 10.0",
            "description": "relevant coursework or notes",
            "achievements": ["coursework items", "academic focus areas"]
        }}
    ],
    "skills": [
        {{"category": "Languages", "items": ["Python", "JavaScript", "..."]}},
        {{"category": "AI / ML", "items": ["LangChain", "..."]}},
        {{"category": "Backend", "items": ["FastAPI", "..."]}},
        {{"category": "Frontend", "items": ["React", "..."]}},
        {{"category": "Databases", "items": ["PostgreSQL", "..."]}},
        {{"category": "DevOps / Cloud", "items": ["Docker", "AWS", "..."]}},
        {{"category": "Tools", "items": ["Git", "..."]}}
    ],
    "projects": [
        {{
            "name": "project name",
            "description": "one-line description",
            "technologies": ["tech1", "tech2"],
            "url": "live url if any",
            "github_url": "github url if any",
            "highlights": ["bullet point 1", "bullet point 2", "...extract ALL bullets"]
        }}
    ],
    "certifications": [
        {{"name": "cert name", "issuer": "issuing org", "date": "", "credential_id": "", "url": ""}}
    ],
    "languages": [
        {{"name": "English", "proficiency": "professional"}},
        {{"name": "Hindi", "proficiency": "native"}}
    ]
}}

IMPORTANT: For skills, group them by meaningful categories (Languages, AI/ML, Backend, Frontend, Databases, DevOps, Tools, etc). Each category has an array of skill names. Do NOT use flat name/level format.
Extract ALL experiences, ALL projects, ALL certifications — do not truncate.

CV Text:
{cv_text}

Return ONLY valid JSON, no explanation."""

    try:
        response = ai_service.client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=8000
        )
        
        result = response.choices[0].message.content.strip()
        
        # Clean up JSON response
        if result.startswith('```json'):
            result = result[7:]
        if result.startswith('```'):
            result = result[3:]
        if result.endswith('```'):
            result = result[:-3]
        
        return json.loads(result)
    except Exception as e:
        logger.error(f"AI parsing error: {str(e)}")
        return None


def merge_cv_data(rule_based: dict, ai_parsed: dict) -> dict:
    """Merge rule-based and AI-parsed CV data, preferring better results"""
    merged = rule_based.copy()
    
    if not ai_parsed:
        return merged
    
    # Merge personal_info - prefer non-empty values
    if 'personal_info' in ai_parsed:
        for key, value in ai_parsed['personal_info'].items():
            if value and (not merged['personal_info'].get(key) or len(str(value)) > len(str(merged['personal_info'].get(key, '')))):
                merged['personal_info'][key] = value
    
    # Prefer AI parsed experience if it has more entries or more details
    if 'experience' in ai_parsed and ai_parsed['experience']:
        ai_exp = ai_parsed['experience']
        rule_exp = merged.get('experience', [])
        
        # Calculate completeness score
        def exp_score(exp_list):
            return sum(
                sum([1 for k in ['title', 'company', 'start_date', 'achievements'] if e.get(k)])
                for e in exp_list
            )
        
        if exp_score(ai_exp) > exp_score(rule_exp):
            merged['experience'] = ai_exp
    
    # Similar for education
    if 'education' in ai_parsed and ai_parsed['education']:
        ai_edu = ai_parsed['education']
        rule_edu = merged.get('education', [])
        
        def edu_score(edu_list):
            return sum(
                sum([1 for k in ['degree', 'institution', 'field_of_study'] if e.get(k)])
                for e in edu_list
            )
        
        if edu_score(ai_edu) > edu_score(rule_edu):
            merged['education'] = ai_edu
    
    # Merge skills — new grouped format {category, items}
    # Prefer AI skills if they have more categories or more total items
    if 'skills' in ai_parsed and ai_parsed['skills']:
        ai_skills = ai_parsed['skills']
        rule_skills = merged.get('skills', [])
        
        def skills_item_count(skill_list):
            return sum(len(s.get('items', [])) for s in skill_list if isinstance(s, dict) and s.get('items'))
        
        if skills_item_count(ai_skills) >= skills_item_count(rule_skills):
            merged['skills'] = ai_skills
        else:
            # Merge: add AI categories not present in rule-based
            existing_cats = {s.get('category', '').lower() for s in rule_skills}
            for sg in ai_skills:
                if sg.get('category', '').lower() not in existing_cats:
                    rule_skills.append(sg)
            merged['skills'] = rule_skills
    
    # Prefer AI for projects if available
    if 'projects' in ai_parsed and ai_parsed['projects'] and len(ai_parsed['projects']) >= len(merged.get('projects', [])):
        merged['projects'] = ai_parsed['projects']
    
    # Merge certifications
    if 'certifications' in ai_parsed and ai_parsed['certifications']:
        merged['certifications'] = ai_parsed['certifications']
    
    # Merge languages
    if 'languages' in ai_parsed and ai_parsed['languages']:
        merged['languages'] = ai_parsed['languages']
    
    return merged


# ============================================
# LaTeX Compilation Endpoint
# ============================================

class CompileLatexRequest(BaseModel):
    latex_source: str
    filename: Optional[str] = "resume"
    enforce_single_page: Optional[bool] = False

def find_pdflatex():
    """Find pdflatex executable on the system"""
    # Check common locations
    paths_to_check = [
        "pdflatex",  # System PATH
    ]
    
    # Windows-specific paths
    if os.name == 'nt':
        paths_to_check.extend([
            r"C:\texlive\2024\bin\windows\pdflatex.exe",
            r"C:\texlive\2023\bin\windows\pdflatex.exe",
            r"C:\Program Files\MiKTeX\miktex\bin\x64\pdflatex.exe",
            r"C:\Users\{}\AppData\Local\Programs\MiKTeX\miktex\bin\x64\pdflatex.exe".format(os.getenv("USERNAME", "")),
        ])
    else:
        paths_to_check.extend([
            "/usr/bin/pdflatex",
            "/usr/local/bin/pdflatex",
            "/usr/local/texlive/2024/bin/x86_64-linux/pdflatex",
        ])
    
    for path in paths_to_check:
        try:
            result = subprocess.run(
                [path, "--version"],
                capture_output=True, timeout=5
            )
            if result.returncode == 0:
                return path
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
            continue
    
    return None


@router.post("/compile-latex")
async def compile_latex(request: CompileLatexRequest):
    """Compile LaTeX source to PDF using pdflatex"""
    pdflatex_path = find_pdflatex()
    
    if not pdflatex_path:
        raise HTTPException(
            status_code=503,
            detail="LaTeX compiler (pdflatex) not available on server. Install TeX Live or MiKTeX."
        )
    
    # Create temporary directory for compilation
    with tempfile.TemporaryDirectory() as tmpdir:
        tex_file = os.path.join(tmpdir, "resume.tex")
        pdf_file = os.path.join(tmpdir, "resume.pdf")
        
        # Write LaTeX source
        with open(tex_file, "w", encoding="utf-8") as f:
            f.write(request.latex_source)
        
        try:
            # Run pdflatex twice for proper cross-references
            for _ in range(2):
                result = subprocess.run(
                    [pdflatex_path, "-interaction=nonstopmode", "-halt-on-error", "resume.tex"],
                    cwd=tmpdir,
                    capture_output=True,
                    text=True,
                    timeout=30
                )
            
            if not os.path.exists(pdf_file):
                # Extract meaningful error from log
                log_file = os.path.join(tmpdir, "resume.log")
                error_msg = "LaTeX compilation failed"
                if os.path.exists(log_file):
                    with open(log_file, "r", encoding="utf-8", errors="ignore") as f:
                        log_content = f.read()
                        # Find error lines
                        errors = [line for line in log_content.split('\n') if line.startswith('!')]
                        if errors:
                            error_msg = "; ".join(errors[:3])
                
                raise HTTPException(status_code=422, detail=error_msg)
            
            # Read the generated PDF
            with open(pdf_file, "rb") as f:
                pdf_bytes = f.read()
            
            # Enforce single-page if requested
            if request.enforce_single_page:
                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    page_count = len(pdf.pages)
                if page_count > 1:
                    raise HTTPException(
                        status_code=422,
                        detail=f"Resume exceeds 1 page ({page_count} pages). Reduce content and try again."
                    )
            
            filename = f"{request.filename or 'resume'}.pdf".replace(" ", "_")
            
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f'inline; filename="{filename}"'
                }
            )
            
        except subprocess.TimeoutExpired:
            raise HTTPException(status_code=504, detail="LaTeX compilation timed out")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"LaTeX compilation error: {e}")
            raise HTTPException(status_code=500, detail=f"Compilation error: {str(e)}")


@router.get("/latex-status")
async def check_latex_status():
    """Check if LaTeX compiler is available"""
    pdflatex_path = find_pdflatex()
    return {
        "available": pdflatex_path is not None,
        "path": pdflatex_path,
        "message": "pdflatex found" if pdflatex_path else "pdflatex not installed. Install TeX Live or MiKTeX for LaTeX compilation."
    }
